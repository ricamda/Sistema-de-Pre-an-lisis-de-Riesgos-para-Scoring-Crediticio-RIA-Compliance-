# app.py - versión final con fix: copiar matriz->riesgos y asegurar 5-7 riesgos
import streamlit as st
import os
import json
import datetime
import logging
import re
from typing import List, Dict, Any

from docs_loader import load_documents_from_folder
from rag import SimpleRAG
from rules import classify_risk_from_inputs
from docx_fill import fill_docx_template

import requests
import pandas as pd

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# ---------- CONFIG ----------
DOCS_FOLDER = "./docs"
TEMPLATE_PATH = "./DATATHON_Modelo Informe_v.0.docx"
OUTPUT_DOCX = "./outputs/informe_generado.docx"
OLLAMA_URL = "http://localhost:11434"
# ----------------------------

st.set_page_config(page_title="Evaluador de Riesgos IA", layout="centered")
st.title("Evaluador preliminar de riesgos — prototipo")
st.markdown("Rellena los campos y pulsa **Analizar**. La herramienta usará los documentos cargados en la carpeta `docs/` como referencias.")

# ---------------- Helpers robustos (JSON parse, sanitización) ----------------

DANGEROUS_KEYWORDS = [
    "matar", "muerte", "disparar", "atacar", "asesinar", "bomb", "suicide", "kill", "shoot", "hurt",
    "explos", "envenen", "tortur"
]
DANGEROUS_REGEX = re.compile(r'\b(' + "|".join(re.escape(k) for k in DANGEROUS_KEYWORDS) + r')\b', flags=re.IGNORECASE)

def _extract_list_items(text: str, max_items: int = 7):
    if not text:
        return []
    items = []
    for line in text.splitlines():
        s = line.strip()
        if re.match(r'^(\-|\•|\*|\d+[\.\)])\s+', s):
            s = re.sub(r'^(\-|\•|\*|\d+[\.\)])\s*', '', s)
            items.append(s)
        else:
            if len(s) > 40:
                items.append(s)
    seen = set()
    out = []
    for it in items:
        if it not in seen:
            out.append(it)
            seen.add(it)
        if len(out) >= max_items:
            break
    return out

def extract_json_from_text(text: str):
    if not text or not text.strip():
        return None
    try:
        return json.loads(text)
    except Exception:
        pass
    first = text.find('{')
    last = text.rfind('}')
    if first != -1 and last != -1 and last > first:
        candidate = text[first:last+1]
        candidate = candidate.replace("“", "\"").replace("”", "\"")
        try:
            return json.loads(candidate)
        except Exception:
            cand2 = re.sub(r'([{,]\s*)([a-zA-Z0-9_]+)\s*:', r'\1"\2":', candidate)
            try:
                return json.loads(cand2)
            except Exception:
                pass
    first_l = text.find('[')
    last_l = text.rfind(']')
    if first_l != -1 and last_l != -1 and last_l > first_l:
        candidate = text[first_l:last_l+1]
        candidate = candidate.replace("“", "\"").replace("”", "\"")
        try:
            return json.loads(candidate)
        except Exception:
            pass
    return None

def sanitize_risk_text(text: str):
    if not text or not text.strip():
        return None, "vacío"
    if DANGEROUS_REGEX.search(text):
        return None, f"Eliminado por contener lenguaje peligroso: '{DANGEROUS_REGEX.search(text).group(0)}'"
    cleaned = " ".join(text.split())
    return cleaned, None

def _parse_probabilidad(val):
    if val is None:
        return None
    try:
        if isinstance(val, (int,float)):
            v = int(val)
            return max(0, min(100, v))
        if isinstance(val, str):
            s = val.strip().replace("%","")
            s = re.sub(r'[^\d\-]', '', s)
            if s == "":
                return None
            v = int(s)
            return max(0, min(100, v))
    except Exception:
        return None
    return None

def normalize_sections_from_llm(text: str):
    sections = {"resumen": "", "riesgos": [], "plan": [], "matriz": [], "fuentes": [], "notes": ""}
    data = extract_json_from_text(text)

    notes = []

    if isinstance(data, dict):
        sections["resumen"] = str(data.get("resumen","") or "")
        raw_riesgos = data.get("riesgos", [])
        if isinstance(raw_riesgos, list):
            for item in raw_riesgos:
                if isinstance(item, dict):
                    desc = item.get("descripcion") or item.get("descripcion_riesgo") or item.get("text") or ""
                else:
                    desc = str(item)
                clean, note = sanitize_risk_text(desc)
                if clean:
                    sections["riesgos"].append(clean)
                elif note:
                    notes.append(note)
        elif isinstance(raw_riesgos, str) and raw_riesgos.strip():
            sections["riesgos"].extend(_extract_list_items(raw_riesgos, max_items=10))
        raw_plan = data.get("plan", [])
        if isinstance(raw_plan, list):
            for it in raw_plan:
                if isinstance(it, dict):
                    desc = it.get("descripcion") or it.get("text") or ""
                    if desc:
                        sections["plan"].append(desc.strip())
                elif isinstance(it, str) and it.strip():
                    sections["plan"].append(it.strip())
        elif isinstance(raw_plan, str):
            sections["plan"].extend(_extract_list_items(raw_plan, max_items=10))
        raw_matriz = data.get("matriz", [])
        if isinstance(raw_matriz, list):
            for row in raw_matriz:
                if not isinstance(row, dict):
                    continue
                desc = row.get("descripcion") or row.get("descripcion_riesgo") or row.get("text") or ""
                clean, note = sanitize_risk_text(desc)
                if not clean:
                    if note: notes.append(note)
                    continue
                prob = _parse_probabilidad(row.get("probabilidad"))
                impacto = row.get("impacto") or row.get("impact") or ""
                sections["matriz"].append({
                    "id": str(row.get("id","")).strip() or "",
                    "descripcion": clean,
                    "probabilidad": prob,
                    "impacto": impacto
                })
        if isinstance(data.get("fuentes"), list):
            sections["fuentes"] = [str(x) for x in data.get("fuentes")]
        if data.get("notes"):
            notes.append(str(data.get("notes")))
        if notes:
            sections["notes"] = "; ".join(notes)
        return sections

    if isinstance(data, list):
        all_strings = all(isinstance(x, str) for x in data)
        all_dicts = all(isinstance(x, dict) for x in data)
        if all_strings:
            for s in data:
                clean, note = sanitize_risk_text(s)
                if clean:
                    sections["riesgos"].append(clean)
                elif note:
                    notes.append(note)
            for i, r in enumerate(sections["riesgos"], start=1):
                sections["matriz"].append({
                    "id": f"R{i:03d}",
                    "descripcion": r,
                    "probabilidad": 80 if i <= 3 else 50,
                    "impacto": "Alto" if "privacidad" in r.lower() or "sesgo" in r.lower() else "Medio"
                })
            sections["notes"] = "; ".join(notes) if notes else ""
            return sections
        elif all_dicts:
            for item in data:
                desc = item.get("descripcion") or item.get("descripcion_riesgo") or item.get("text") or ""
                clean, note = sanitize_risk_text(desc)
                if not clean:
                    if note: notes.append(note)
                    continue
                prob = _parse_probabilidad(item.get("probabilidad"))
                impacto = item.get("impacto") or item.get("impact") or ""
                rid = str(item.get("id","")).strip()
                if not rid:
                    rid = f"R{len(sections['matriz'])+1:03d}"
                sections["matriz"].append({
                    "id": rid,
                    "descripcion": clean,
                    "probabilidad": prob,
                    "impacto": impacto
                })
                sections["riesgos"].append(clean)
            sections["notes"] = "; ".join(notes) if notes else ""
            return sections
        else:
            for it in data:
                if isinstance(it, str):
                    clean, note = sanitize_risk_text(it)
                    if clean:
                        sections["riesgos"].append(clean)
                    elif note:
                        notes.append(note)
                elif isinstance(it, dict):
                    desc = it.get("descripcion") or it.get("text") or ""
                    clean, note = sanitize_risk_text(desc)
                    if clean:
                        sections["riesgos"].append(clean)
                        prob = _parse_probabilidad(it.get("probabilidad"))
                        sections["matriz"].append({
                            "id": str(it.get("id","")).strip() or f"R{len(sections['matriz'])+1:03d}",
                            "descripcion": clean,
                            "probabilidad": prob,
                            "impacto": it.get("impacto","")
                        })
                    elif note:
                        notes.append(note)
            sections["notes"] = "; ".join(notes) if notes else ""
            return sections

    parts = [p.strip() for p in text.split("\n\n") if p.strip()]
    if parts:
        sections["resumen"] = parts[0]
    riesgos_candidates = []
    plan_candidates = []
    for p in parts[1:]:
        low = p.lower()
        if any(k in low for k in ["riesgo", "riesgos", "riesgos principales"]):
            riesgos_candidates.extend(_extract_list_items(p, max_items=10))
        elif any(k in low for k in ["plan de acción", "plan de accion", "plan", "medida", "acciones", "mitig"]):
            plan_candidates.extend(_extract_list_items(p, max_items=10))
    if not riesgos_candidates:
        riesgos_candidates = _extract_list_items(text, max_items=7)
    if not plan_candidates:
        plan_candidates = _extract_list_items(text, max_items=7)
    for r in riesgos_candidates[:7]:
        clean, note = sanitize_risk_text(r)
        if clean:
            sections["riesgos"].append(clean)
        elif note:
            notes.append(note)
    sections["plan"] = plan_candidates[:7]
    for idx, r in enumerate(sections["riesgos"], start=1):
        sections["matriz"].append({
            "id": f"R{idx:03d}",
            "descripcion": r,
            "probabilidad": 80 if idx <= 3 else 50,
            "impacto": "Alto" if "privacidad" in r.lower() or "sesgo" in r.lower() else "Medio"
        })
    sections["notes"] = "; ".join(notes) if notes else ""
    return sections

# ---------- Nuevas utilidades: asegurar 5-7 riesgos y asignar etiquetas -----------
def ensure_risks_count(risks: List[str], nivel_riesgo: str, target_min: int = 5, target_max: int = 7) -> List[str]:
    out = list(dict.fromkeys(risks))
    if len(out) >= target_min:
        return out[:target_max]
    fallbacks = []
    if nivel_riesgo.lower().startswith("alto"):
        fallbacks = [
            "Riesgo de privacidad por tratamiento de datos sensibles",
            "Riesgo de discriminación o sesgo en scoring",
            "Riesgo de decisiones incorrectas por falta de supervisión humana",
            "Riesgo de fuga o exposición de datos a terceros",
            "Riesgo de incumplimiento regulatorio por falta de documentación",
            "Riesgo operativo por vulnerabilidades en la infraestructura"
        ]
    else:
        fallbacks = [
            "Riesgo de privacidad",
            "Riesgo operativo",
            "Riesgo de transparencia",
            "Riesgo de disponibilidad"
        ]
    for f in fallbacks:
        if len(out) >= target_min:
            break
        if f not in out:
            out.append(f)
    i = 1
    while len(out) < target_min:
        candidate = f"Riesgo adicional no especificado {i}"
        out.append(candidate)
        i += 1
    return out[:target_max]

def label_probability_from_score(score: int) -> str:
    if score is None:
        return "Media"
    if score >= 75:
        return "Alta"
    if score >= 40:
        return "Media"
    return "Baja"

def label_impact_from_text(text: str, default="Medio") -> str:
    low = text.lower()
    if any(k in low for k in ["gran", "alto", "grave", "severa", "significativo", "significativa"]):
        return "Alto"
    if any(k in low for k in ["medio", "moderado", "intermedio"]):
        return "Medio"
    if any(k in low for k in ["bajo", "mínimo", "minimo"]):
        return "Bajo"
    return default

def infer_probability_score_from_keywords(text: str) -> int:
    low = text.lower()
    if any(k in low for k in ["muy probable", "alta probabilidad", "probable", "alta"]):
        return 80
    if any(k in low for k in ["posible", "media", "moderada"]):
        return 50
    if any(k in low for k in ["raro", "baja probabilidad", "bajo"]):
        return 20
    return 50

def infer_tipologia(text: str):
    low = text.lower()
    if any(k in low for k in ["privacidad", "datos", "gdpr", "protección", "fuga"]):
        return "Normativo"
    if any(k in low for k in ["sesgo", "bias", "equidad", "fairness", "discrimin"]):
        return "Técnico"
    if any(k in low for k in ["seguridad", "vulnerab", "ataque", "segur"]):
        return "Operativo"
    if any(k in low for k in ["transpar", "explic", "interpret"]):
        return "Normativo"
    return "Técnico"

def enrich_risks(risks_list: List[str], nivel_riesgo: str, rol_entidad: str = "Proveedor"):
    enriched = []
    risks_list = ensure_risks_count(risks_list, nivel_riesgo, target_min=5, target_max=7)
    for i, r in enumerate(risks_list, start=1):
        tip = infer_tipologia(r := r)
        score = infer_probability_score_from_keywords(r)
        score = min(95, score + max(0, 15 - (i-1)*3))
        prob_label = label_probability_from_score(score)
        impact_label = label_impact_from_text(r, default=("Alto" if nivel_riesgo.lower().startswith("alto") else "Medio"))
        if nivel_riesgo.lower().startswith("alto"):
            articulo = "RIA Art. 6 — Evaluación y obligaciones adicionales para sistemas de alto riesgo (proveedor)."
        elif "limit" in nivel_riesgo.lower():
            articulo = "RIA Art. 5 — Requisitos de transparencia y gobernanza (riesgo limitado)."
        else:
            articulo = "RIA Art. 4 — Buenas prácticas y recomendaciones (riesgo mínimo)."
        enriched.append({
            "id": f"R{i:03d}",
            "descripcion": r,
            "tipologia": tip,
            "probabilidad": prob_label,
            "impacto": impact_label,
            "prob_score": score,
            "articulo_ria": articulo
        })
    return enriched

# ----------------- Plan: asignar urgencias según riesgos mitigados ------------
def assign_urgency_to_measures(measures: List[str], enriched_risks: List[Dict[str,Any]]):
    out = []
    for m in measures:
        desc = m.strip()
        low = desc.lower()
        urg = "Media"
        if any(k in low for k in ["supervisión", "humana", "apelar", "appeal", "revisión"]):
            urg = "Alta"
        if any(k in low for k in ["seguridad", "cifrar", "encript", "auth", "autentic", "mfa", "acceso"]):
            urg = "Alta"
        if any(k in low for k in ["auditor", "auditoría", "auditoria", "certificación"]):
            urg = "Media"
        if any(k in low for k in ["minimizar", "logging", "registro", "trazabilidad"]):
            urg = "Media"
        for r in enriched_risks:
            if r["id"].lower() in low or any(w in low for w in r["descripcion"].lower().split()[:3]):
                if r.get("probabilidad") == "Alta" or r.get("impacto") == "Alto":
                    urg = "Alta"
                elif r.get("probabilidad") == "Media":
                    urg = max(urg, "Media")
        out.append({"nombre": desc, "urgencia": urg})
    return out

# ---------------- build_replacements_v2 (añade urgencias) ----------------
def build_replacements_v2(tpl_placeholders: List[str], sections: Dict[str,Any], nombre: str, finalidad: str,
                          nivel: str, fecha_iso: str, rule_res: Dict[str,Any], sources: List[Dict]=None,
                          max_riesgos: int = 7, max_medidas: int = 7):
    if sources is None:
        sources = []
    rp = {}
    rp["{{nombre_caso_uso}}"] = nombre
    rp["{{descripcion_ejecutiva}}"] = finalidad
    rp["{{resumen_ejecutivo}}"] = sections.get("resumen","")
    rp["{{fecha_analisis}}"] = fecha_iso
    rp["{{nivel_riesgo}}"] = nivel
    rp["{{clasificacion_regulatoria}}"] = nivel
    rp["{{justificacion_clasificacion}}"] = rule_res.get("justificacion","")
    rp["{{anexo_fuentes}}"] = sections.get("fuentes","") or ("Fuentes recuperadas: " + ", ".join([s["id"] for s in sources[:6]]))

    riesgos_llm = sections.get("riesgos", []) or []
    medidas_llm = sections.get("plan", []) or []

    riesgos_llm = ensure_risks_count(riesgos_llm, nivel, target_min=5, target_max=7)
    if not medidas_llm:
        medidas_llm = [
            "Implementar supervisión humana en decisiones críticas y puntos de apelación.",
            "Realizar auditorías de sesgo y fairness periódicas.",
            "Minimizar y cifrar datos sensibles en pipelines de entrenamiento y predicción.",
            "Mantener logging de decisiones y razonamiento para trazabilidad.",
            "Establecer cláusulas contractuales y controles de acceso con proveedores."
        ][:max_medidas]
    medidas_llm = medidas_llm[:max_medidas]

    enriched = enrich_risks(riesgos_llm, nivel, rol_entidad=sections.get("rol_entidad","Proveedor"))

    for i, r in enumerate(enriched, start=1):
        rp[f"{{{{id_riesgo_{i}}}}}"] = r["id"]
        rp[f"{{{{descripcion_riesgo_{i}}}}}"] = r["descripcion"]
        rp[f"{{{{tipologia_riesgo_{i}}}}}"] = r["tipologia"]
        rp[f"{{{{probabilidad_{i}}}}}"] = r["probabilidad"]
        rp[f"{{{{impacto_{i}}}}}"] = r["impacto"]
        rp[f"{{{{nivel_riesgo_matriz_{i}}}}}"] = nivel
        rp[f"{{{{articulo_ria_{i}}}}}"] = r["articulo_ria"]

    measures_with_urgency = assign_urgency_to_measures(medidas_llm, enriched)
    for j, m in enumerate(measures_with_urgency, start=1):
        rp[f"{{{{nombre_medida_{j}}}}}"] = m["nombre"]
        rp[f"{{{{descripcion_medida_{j}}}}}"] = m["nombre"]
        rp[f"{{{{urgencia_{j}}}}}"] = m["urgencia"]
        rp[f"{{{{responsable_{j}}}}}"] = "Técnico"

    rp["{{tabla_matriz_riesgos}}"] = "__INSERT_MATRIZ_RIESGOS__"
    rp["{{tabla_plan_accion}}"] = "__INSERT_PLAN_ACCION__"

    for ph in tpl_placeholders:
        key = f"{{{{{ph}}}}}"
        if key in rp:
            continue
        lower = ph.lower()
        if "nombre" in lower:
            rp[key] = nombre
            continue
        if "fecha" in lower:
            rp[key] = fecha_iso
            continue
        if "resumen" in lower:
            rp[key] = rp.get("{{resumen_ejecutivo}}", "")
            continue
        if "justificacion" in lower:
            rp[key] = rp.get("{{justificacion_clasificacion}}", "")
            continue
        if "anexo" in lower or "fuente" in lower:
            rp[key] = rp.get("{{anexo_fuentes}}", "")
            continue
        rp[key] = f"[REVISAR: completar '{ph}']"
    return rp

# ----------------- Cargar docs y RAG -----------------
@st.cache_data(show_spinner=False)
def load_and_index_docs(folder: str):
    docs = load_documents_from_folder(folder)
    return docs

docs = load_and_index_docs(DOCS_FOLDER)
if not docs:
    st.warning(f"No se encontraron documentos en '{DOCS_FOLDER}'. Añade los PDFs/DOCX/XLSX allí y vuelve a ejecutar.")
    st.stop()

rag = SimpleRAG(docs)

# ----------------- Formulario -----------------
with st.form("form"):
    nombre = st.text_input("Nombre del sistema", value="CreditAI")
    finalidad = st.text_area("Finalidad prevista (qué hace la IA)", value="Evaluar riesgo crediticio")
    sector = st.selectbox("Sector", ["Financiero","Salud","RRHH","Otro"])
    tipo_decision = st.selectbox("Tipo de decisión", ["Totalmente automatizada","Asistida por humano","Solo soporte"])
    datos_sensibles = st.selectbox("Procesa datos sensibles?", ["si","no"])
    impacto = st.selectbox("Impacto de errores", ["Alto","Medio","Bajo"])
    rol_entidad = st.selectbox("Rol de la entidad", ["Proveedor","Importador","Responsable","Operador"])
    submit = st.form_submit_button("Analizar")

if not submit:
    st.info("Completa el formulario y pulsa Analizar.")
    st.stop()

st.info("Ejecutando análisis...")

inputs = {"sector": sector, "tipo_decision": tipo_decision, "datos_sensibles": datos_sensibles, "impacto": impacto, "finalidad": finalidad, "rol_entidad": rol_entidad}

# 1) Motor de reglas
rule_res = classify_risk_from_inputs(inputs)
nivel = rule_res.get("nivel_riesgo", "Riesgo desconocido")
justificacion = rule_res.get("justificacion", "")
score = rule_res.get("score", None)

st.subheader("Resultado del motor de reglas")
st.write(f"**Nivel:** {nivel}")
st.write(f"**Puntuación interna:** {score}")
st.write(f"**Justificación:** {justificacion}")

# 2) RAG y context
user_query = f"Sistema: {nombre}. Finalidad: {finalidad}. Sector: {sector}. Tipo decisión: {tipo_decision}. Datos sensibles: {datos_sensibles}. Impacto: {impacto}."
context_str, sources = rag.build_context(user_query, top_k=6, token_budget_chars=1400)

st.subheader("Evidencia recuperada (RAG)")
if sources:
    for s in sources[:6]:
        meta = s.get("meta", {})
        st.markdown(f"- **{s['id']}** (score={s['score']:.3f}) — fuente: {meta.get('source','-')} page:{meta.get('page','-')}")
else:
    st.write("No se encontraron fragmentos relevantes.")

# 3) Prompt (pedimos JSON)
system_prompt = (
    "Eres un auditor experto en regulación de IA. Usa SOLO la información provista en CONTEXT y INPUT. "
    "RESPONDE EXCLUSIVAMENTE en JSON válido con el siguiente esquema:\n"
    '{"resumen":"", "riesgos":["..."], "plan":["..."], "matriz":[{"id":"R001","descripcion":"","probabilidad":"","impacto":""}], "fuentes":["..."], "notes":""}\n'
    "Si no puedes completar un campo con evidencia, deja la cadena vacía o la lista vacía. No escribas texto fuera del JSON."
)
prompt = f"{system_prompt}\n\nCONTEXT:\n{context_str}\n\nINPUT:\n{user_query}\n\nRESPONDE:"

# 4) Llamada LLM
llm_text = ""
try:
    payload = {"model": "nemotron-mini", "prompt": prompt, "stream": False, "max_tokens": 1200}
    resp = requests.post(f"{OLLAMA_URL}/api/generate", json=payload, timeout=120)
    if resp.status_code != 200:
        st.error(f"Error llamando a Ollama: {resp.status_code} - {resp.text[:1000]}")
        st.stop()
    data = resp.json()
    if isinstance(data, dict):
        llm_text = data.get("response") or data.get("text") or data.get("output") or ""
        if not llm_text and "choices" in data:
            choices = data["choices"]
            if isinstance(choices, list) and choices:
                first = choices[0]
                llm_text = first.get("text") or first.get("message", {}).get("content", "") or ""
    if not llm_text:
        llm_text = resp.text
except Exception as e:
    st.error(f"No se pudo conectar con Ollama: {e}")
    st.stop()

st.subheader("Salida LLM (borrador)")
st.text_area("Salida LLM", value=llm_text, height=300)

# 5) Parseo LLM -> secciones
sections = normalize_sections_from_llm(llm_text)
# ------------------ PARCHE: normalizar riesgos y crear plan con urgencia ------------------

def _is_embedded_json_str(s: str) -> bool:
    if not isinstance(s, str):
        return False
    s2 = s.strip()
    return any(k in s2 for k in ['"resumen"', '"riesgos"', '"matriz"', '"plan"', '"descripcion"']) and len(s2) > 30

def _clean_text(s):
    if not isinstance(s, str):
        return ""
    # quitar comillas sobrantes y saltos múltiples
    return s.strip().strip('"').strip()

def infer_urgency_from_prob_impact(prob_label: str, impact_label: str) -> str:
    """
    Heurística simple:
    - Si probabilidad Alta o impacto Alto => Alta
    - Si (probabilidad Media y impacto Medio) => Media
    - Si ambos bajos => Baja
    - Otros casos razonables => Media
    """
    prob = (prob_label or "").lower()
    impact = (impact_label or "").lower()
    if "alta" in prob or "alto" in impact:
        return "Alta"
    if "baja" in prob and "bajo" in impact:
        return "Baja"
    # caso por defecto
    return "Media"

def generate_mitigation_for_risk(risk_text: str) -> str:
    """
    Mapeo heurístico de riesgos a medidas concretas (plantillas cortas).
    Añade una medida principal orientada al tipo de riesgo.
    """
    rt = (risk_text or "").lower()
    if any(w in rt for w in ["privacidad", "datos sensibles", "protección de datos", "gdpr", "aepd"]):
        return "Minimizar y cifrar datos sensibles; aplicar acceso por roles y anonimización cuando sea posible."
    if any(w in rt for w in ["sesgo", "discrimin", "bias", "fairness"]):
        return "Realizar evaluación de sesgo y fairness, re-equilibrado de datos y controles de métricas de equidad."
    if any(w in rt for w in ["supervisión", "humana", "apelar", "recurso"]):
        return "Implementar supervisión humana obligatoria en decisiones críticas y un proceso de apelación."
    if any(w in rt for w in ["seguridad", "vulnerab", "exposición", "fuga"]):
        return "Aplicar controles de seguridad: MFA, segmentación, monitorización y pruebas de penetración."
    if any(w in rt for w in ["transparencia", "explicab", "interpretab", "documentac"]):
        return "Documentar modelos, explicar características relevantes y publicar métricas de desempeño y límites."
    if any(w in rt for w in ["operativo", "disponibilidad", "infraestruct"]):
        return "Establecer SRE/operaciones, recuperación ante desastres y monitorización de SLAs."
    # fallback general
    return "Definir políticas de gobernanza, realizar auditorías periódicas y controles técnicos y organizativos."

# 1) Limpiar riesgos detectados
raw_riesgos = sections.get("riesgos") or []
clean_riesgos = []
for r in raw_riesgos:
    if isinstance(r, dict):
        desc = r.get("descripcion") or r.get("descripcion_riesgo") or r.get("text") or ""
    else:
        desc = r
    desc = _clean_text(desc)
    if desc and not _is_embedded_json_str(desc):
        clean_riesgos.append(desc)
# deduplicate preserving order
seen = set()
clean_riesgos = [x for x in clean_riesgos if not (x in seen or seen.add(x))]

sections["riesgos"] = clean_riesgos

# 2) Normalizar matriz si existe (extraer probabilidad/impacto) -> nos sirve para urgencias
# construir diccionario de metadata por riesgo (por texto)
risk_meta = {}
for m in (sections.get("matriz") or []):
    if isinstance(m, dict):
        key = _clean_text(m.get("descripcion","") or m.get("descripcion_riesgo","") or "")
        if not key:
            continue
        prob = m.get("probabilidad") or m.get("prob","") or ""
        impact = m.get("impacto") or m.get("impact","") or ""
        risk_meta[key] = {"probabilidad": _clean_text(prob), "impacto": _clean_text(impact)}

# 3) Limpiar y normalizar plan devuelto por LLM
raw_plan = sections.get("plan") or []
clean_plan = []
for p in raw_plan:
    if isinstance(p, dict):
        # puede venir como {id, descripcion, urgencia}
        desc = p.get("descripcion") or p.get("medida") or p.get("text") or ""
    else:
        desc = p
    desc = _clean_text(desc)
    if not desc or _is_embedded_json_str(desc):
        continue
    clean_plan.append(desc)
# deduplicate
seen = set()
clean_plan = [x for x in clean_plan if not (x in seen or seen.add(x))]

# 4) Si el plan devuelto son básicamente repeticiones (ej, sus items aparecen dentro de riesgos),
# generar medidas útiles por cada riesgo usando la función heurística
# criterio: si no hay al menos N medidas originales o las medidas son descriptivas (contienen texto similar a los riesgos),
# entonces generamos medidas por riesgo.
def _is_plan_simply_risks(plan_list, risks_list):
    if not plan_list:
        return True
    # si más del 50% de medidas son substring de algún riesgo -> es repetitivo
    count_sim = 0
    for pm in plan_list:
        for rk in risks_list:
            if pm.lower() in rk.lower() or rk.lower() in pm.lower():
                count_sim += 1
                break
    return (count_sim / max(1, len(plan_list))) > 0.5

if _is_plan_simply_risks(clean_plan, sections["riesgos"]) or len(clean_plan) < max(3, len(sections["riesgos"])//2):
    # generamos medidas a partir de riesgos
    generated_plan = []
    for r in sections["riesgos"]:
        gen = generate_mitigation_for_risk(r)
        if gen:
            generated_plan.append(gen)
    # si hay menos medidas que riesgos, pad con medidas generales
    if len(generated_plan) < 5:
        extras = [
            "Realizar auditoría independiente periódica del sistema.",
            "Establecer controles contractuales y gestión de proveedores.",
            "Implementar registro detallado (logging) de decisiones y explicaciones."
        ]
        for e in extras:
            if e not in generated_plan:
                generated_plan.append(e)
    clean_plan = generated_plan

# 5) Asignar urgencias a cada medida basadas en la matriz (si existe) o por heurística de riesgo
measures_with_urgency = []
for idx, meas in enumerate(clean_plan, start=1):
    # intento asociar medida a riesgo relevante buscando palabras compartidas
    assoc_prob = ""
    assoc_impact = ""
    for rk_text, meta in risk_meta.items():
        # si la medida menciona palabras del riesgo asociamos
        if any(w in meas.lower() for w in rk_text.lower().split()[:4]):
            assoc_prob = meta.get("probabilidad","")
            assoc_impact = meta.get("impacto","")
            break
    # si no encontramos asociación, intentar usar la peor prob/impact de la matriz si existe
    if not assoc_prob and risk_meta:
        # tomar la primera meta disponible
        some = next(iter(risk_meta.values()))
        assoc_prob = some.get("probabilidad","")
        assoc_impact = some.get("impacto","")
    urgency = infer_urgency_from_prob_impact(assoc_prob, assoc_impact)
    measures_with_urgency.append({"medida": meas, "urgencia": urgency, "id": f"M{idx:02d}"})

# 6) Guardar resultados limpios en sections para uso posterior (UI / llenado docx)
sections["plan"] = [m["medida"] for m in measures_with_urgency]
sections["plan_struct"] = measures_with_urgency  # estructura con urgencia e id
# 7) Reconstruir la 'matriz' si viene vacía para asegurar coherencia entre riesgos y tabla
if not sections.get("matriz") or len(sections.get("matriz")) == 0:
    nueva_matriz = []
    for i, desc in enumerate(sections.get("riesgos", []), start=1):
        meta = risk_meta.get(desc, {})
        prob_label = meta.get("probabilidad") or ("Alta" if "alto" in desc.lower() else "Media")
        impact_label = meta.get("impacto") or ("Alto" if "alto" in desc.lower() else "Medio")
        nueva_matriz.append({
            "id": f"R{i:03d}",
            "descripcion": desc,
            "probabilidad": _clean_text(prob_label),
            "impacto": _clean_text(impact_label),
            "tipologia": "Normativo / Técnico / Operativo"
        })
    sections["matriz"] = nueva_matriz

# fin del parche de normalización
# ------------------------------------------------------------------------------------
sections["rol_entidad"] = rol_entidad

# --- FIX: si la LLM devolvió la info en 'matriz' pero no en 'riesgos', trasladarla ---
if (not sections.get("riesgos")) and sections.get("matriz"):
    try:
        sections["riesgos"] = [m.get("descripcion") for m in sections.get("matriz", []) if m.get("descripcion")]
    except Exception:
        sections["riesgos"] = []

# --- FIX 2: asegurar entre 5 y 7 riesgos antes de enriquecer (fallbacks) ---
sections["riesgos"] = ensure_risks_count(sections.get("riesgos", []), nivel, target_min=5, target_max=7)

# 6) Fallbacks y enriquecimiento
if not sections.get("plan"):
    sections["plan"] = []

enriched = enrich_risks(sections["riesgos"], nivel, rol_entidad)

# 7) Mostrar UI: riesgos (5-7) y plan con urgencia
st.subheader("Riesgos principales (enriquecidos)")
if enriched:
    df_rows = []
    for r in enriched:
        df_rows.append({
            "ID": r["id"],
            "Riesgo (descripción)": r["descripcion"],
            "Tipología": r["tipologia"],
            "Probabilidad": r["probabilidad"],
            "Impacto": r["impacto"],
            "Artículo RIA (aplicable)": r["articulo_ria"]
        })
    df = pd.DataFrame(df_rows)
    st.dataframe(df, use_container_width=True)
else:
    st.write("No se han identificado riesgos.")

# asegurar que el plan tenga al menos 5 medidas
if not sections.get("plan"):
    sections["plan"] = [
        "Implementar supervisión humana en decisiones críticas y puntos de apelación.",
        "Realizar auditorías de sesgo y fairness periódicas.",
        "Minimizar y cifrar datos sensibles en pipelines de entrenamiento y predicción.",
        "Mantener logging de decisiones y razonamiento para trazabilidad.",
        "Establecer cláusulas contractuales y controles de acceso con proveedores."
    ]
measures_with_urgency = assign_urgency_to_measures(sections["plan"], enriched)

st.subheader("Plan de acción (propuesto) — con urgencia")
for i, m in enumerate(measures_with_urgency, start=1):
    st.markdown(f"{i}. {m['nombre']} — **Urgencia:** {m['urgencia']}")

# Matriz breve
st.subheader("Matriz breve (resumen por riesgo)")
for r in enriched:
    st.markdown(f"- **{r['id']}** — {r['descripcion']}")
    st.markdown(f"    - Tipología: **{r['tipologia']}**, Probabilidad: **{r['probabilidad']}**, Impacto: **{r['impacto']}**")
    st.markdown(f"    - Artículo RIA sugerido: _{r['articulo_ria']}_")

# 8) Preparar replacements y guardar debug
fecha_iso = datetime.datetime.utcnow().strftime("%Y-%m-%d")
def list_placeholders(path):
    from docx import Document
    doc = Document(path)
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    full = "\n".join(texts)
    placeholders = []
    try:
        placeholders = re.findall(r"\{\{([^}]+)\}\}", full)
    except Exception:
        placeholders = []
    return sorted(set(placeholders))

tpl = TEMPLATE_PATH if os.path.exists(TEMPLATE_PATH) else None
tpl_placeholders = []
if tpl:
    try:
        tpl_placeholders = list_placeholders(tpl)
    except Exception:
        tpl_placeholders = []

replacements = build_replacements_v2(tpl_placeholders, {"resumen": sections.get("resumen",""), "riesgos": [r["descripcion"] for r in enriched], "plan": [m["nombre"] for m in measures_with_urgency], "fuentes": sections.get("fuentes",[]), "rol_entidad": rol_entidad}, nombre, finalidad, nivel, fecha_iso, rule_res, sources=sources)

os.makedirs("./outputs", exist_ok=True)
debug_obj = {
    "inputs": inputs,
    "rule_res": rule_res,
    "rag_sources": sources,
    "prompt": prompt,
    "llm_text": llm_text[:20000],
    "sections": sections,
    "enriched_risks": enriched,
    "measures_with_urgency": measures_with_urgency,
    "replacements_preview": {k: replacements[k] for k in list(replacements)[:40]}
}
debug_path = f"./outputs/debug_{nombre.replace(' ','_')}_{fecha_iso}.json"
with open(debug_path, "w", encoding="utf-8") as f:
    json.dump(debug_obj, f, ensure_ascii=False, indent=2)

st.info(f"Debug guardado en {debug_path}")

# 9) Generar DOCX (opcional)
if tpl:
    try:
        os.makedirs(os.path.dirname(OUTPUT_DOCX) or ".", exist_ok=True)
        fill_docx_template(tpl, OUTPUT_DOCX, replacements)
        with open(OUTPUT_DOCX, "rb") as f:
            doc_bytes = f.read()
        st.success("Documento DOCX generado (opcional).")
        st.download_button("Descargar informe DOCX", data=doc_bytes, file_name=os.path.basename(OUTPUT_DOCX),
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        st.warning(f"No se pudo generar DOCX: {e}")
else:
    st.info("Plantilla DOCX no encontrada; se omitió la generación de DOCX (la interfaz contiene toda la información).")

# FIN